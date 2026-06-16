"""
export_manuscript.py
====================

Exportiert fertige DE-Szenen als DOCX und/oder EPUB.

Quelle:
    books/<book-id>/work/scenes/de/<style>/<Kapitel>/scene-XX.md

Ausgabe:
    books/<book-id>/exports/<style>/chapter/docx/
    books/<book-id>/exports/<style>/chapter/epub/
    books/<book-id>/exports/<style>/chapter/work/
    books/<book-id>/exports/<style>/book/docx/
    books/<book-id>/exports/<style>/book/epub/
    books/<book-id>/exports/<style>/book/work/
"""

from __future__ import annotations

import argparse
import html
import io
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).parent))

import yaml

from lib.book_project import find_book as find_book_project
from lib.output_paths import (
    book_exports_root,
    book_output_root,
    find_scene_translations,
    list_chapter_ids_with_source_scenes,
    list_source_scene_paths,
    parse_scene_number,
    source_chapter_path,
)


REPO_ROOT = Path(__file__).resolve().parent.parent
EXPORT_FORMATS = ("docx", "epub", "pdf", "all")


@dataclass
class SceneExport:
    number: int
    text: str


@dataclass
class ChapterExport:
    chapter_id: str
    title: str
    scenes: list[SceneExport]
    missing: list[int]
    ru_count: int
    de_count: int


@dataclass
class ExportResult:
    chapters: list[ChapterExport]
    missing_by_chapter: dict[str, list[int]]
    partial: bool


@dataclass(frozen=True)
class Illustration:
    kind: str
    chapter_id: str
    scene_number: int | None
    path: Path


def load_yaml(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    data = yaml.safe_load(path.read_text(encoding="utf-8"))
    return data or {}


def find_book(book_id: str | None) -> dict[str, Any]:
    return find_book_project(REPO_ROOT, book_id)


def load_export_config(book: dict[str, Any]) -> dict[str, Any]:
    export_path = REPO_ROOT / book.get("export_config", "")
    data = load_yaml(export_path)
    defaults = data.get("defaults", {}) or {}
    book_cfg = data.get("book", {}) or {}
    meta = {**defaults, **book_cfg}
    for key in ("cover", "front_matter", "output", "illustrations"):
        merged = {
            **(defaults.get(key, {}) or {}),
            **(book_cfg.get(key, {}) or {}),
        }
        if merged:
            meta[key] = merged
    meta.setdefault("title", book.get("title", ""))
    meta.setdefault("author", book.get("author", ""))
    meta.setdefault("language", "de-DE")
    meta.setdefault("cover", defaults.get("cover", {}) or {})
    meta.setdefault("front_matter", defaults.get("front_matter", {}) or {})
    meta.setdefault("output", defaults.get("output", {}) or {})
    meta.setdefault("illustrations", defaults.get("illustrations", {}) or {})
    structure = book.get("structure") or {}
    meta.setdefault("structure_groups", structure.get("groups") or [])
    meta.setdefault("display", book.get("display") or {})
    meta["_base_dir"] = str(export_path.parent)
    return meta


def sanitize_filename(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    value = value.strip("-")
    return value or "export"


def yaml_scalar(value: Any) -> str:
    dumped = yaml.safe_dump(
        str(value),
        default_flow_style=True,
        allow_unicode=True,
        sort_keys=False,
    )
    return dumped.splitlines()[0]


def export_dirs(exports_root: Path, style: str, scope: str) -> dict[str, Path]:
    root = exports_root / style / scope
    return {
        "root": root,
        "docx": root / "docx",
        "epub": root / "epub",
        "pdf": root / "pdf",
        "work": root / "work",
        "manifests": root / "manifests",
    }


def get_title(output_root: Path, chapter_id: str) -> str:
    src = source_chapter_path(output_root, chapter_id)
    if not src.exists():
        return f"Kapitel {chapter_id}"
    for line in src.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            return title or f"Kapitel {chapter_id}"
    return f"Kapitel {chapter_id}"


CONTROL_LINE_PATTERNS = [
    # Scene files may contain LLM-generated wrappers like "## 6",
    # "### Vier" or "## Szene 4". Reader exports should expose only
    # chapter-level structure; Pandoc uses Markdown headings for EPUB nav.
    re.compile(r"^#{1,6}\s+.+$"),
    re.compile(r"^\d+\s*$"),
    re.compile(r"^#{1,6}\s*szene\s+\d+\s*$", re.IGNORECASE),
    re.compile(r"^#{1,6}\s*scene\s+\d+\s*$", re.IGNORECASE),
    re.compile(r"^\*?buch:\s+.*\*?$", re.IGNORECASE),
    re.compile(r"^\*?stil:\s+.*\*?$", re.IGNORECASE),
    re.compile(r"^\*?erstellt am:\s+.*\*?$", re.IGNORECASE),
    re.compile(r"^\*?provider:\s+.*\*?$", re.IGNORECASE),
    re.compile(r"^\*?modell:\s+.*\*?$", re.IGNORECASE),
    re.compile(r"^\*?tokens?:\s+.*\*?$", re.IGNORECASE),
    re.compile(r"^-{3,}$"),
]


def clean_scene_markdown(text: str) -> str:
    lines = []
    for line in text.replace("\r\n", "\n").split("\n"):
        stripped = line.strip()
        if any(pattern.match(stripped) for pattern in CONTROL_LINE_PATTERNS):
            continue
        if stripped.startswith("Hier ist die ") or stripped.startswith(
            "Hier ist eine "
        ):
            continue
        lines.append(line.rstrip())
    cleaned = "\n".join(lines).strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


def markdown_to_plain_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for raw_block in re.split(r"\n\s*\n", text.strip()):
        block = raw_block.strip()
        if not block:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", block)
        if heading:
            level = min(len(heading.group(1)), 3)
            blocks.append((f"heading{level}", heading.group(2).strip()))
        else:
            block = re.sub(r"\*\*(.*?)\*\*", r"\1", block)
            block = re.sub(r"\*(.*?)\*", r"\1", block)
            blocks.append(("paragraph", block.replace("\n", " ")))
    return blocks


def collect_chapter(
    output_root: Path,
    chapter_id: str,
    style: str,
    allow_partial: bool,
    source_lang: str = "ru",
) -> ChapterExport:
    source_paths = list_source_scene_paths(output_root, chapter_id, source_lang)
    source_nums = [
        num for path in source_paths
        if (num := parse_scene_number(path, chapter_id)) is not None
    ]
    scene_map = find_scene_translations(output_root, chapter_id, style)
    missing = [num for num in sorted(source_nums) if num not in scene_map]
    if missing and not allow_partial:
        return ChapterExport(
            chapter_id=chapter_id,
            title=get_title(output_root, chapter_id),
            scenes=[],
            missing=missing,
            ru_count=len(source_nums),
            de_count=len(scene_map),
        )
    scenes = []
    for num in sorted(scene_map):
        if source_nums and num not in set(source_nums):
            continue
        text = clean_scene_markdown(scene_map[num].read_text(encoding="utf-8"))
        if text:
            scenes.append(SceneExport(number=num, text=text))
    return ChapterExport(
        chapter_id=chapter_id,
        title=get_title(output_root, chapter_id),
        scenes=scenes,
        missing=missing,
        ru_count=len(source_nums),
        de_count=len(scene_map),
    )


def collect_export(
    output_root: Path,
    style: str,
    scope: str,
    chapter_id: str | None,
    allow_partial: bool,
    source_lang: str = "ru",
) -> ExportResult:
    if scope == "chapter":
        if not chapter_id:
            raise SystemExit("--chapter ist bei --scope chapter erforderlich")
        chapter_ids = [chapter_id]
    else:
        chapter_ids = list_chapter_ids_with_source_scenes(output_root, source_lang)
    chapters = [
        collect_chapter(output_root, cid, style, allow_partial, source_lang)
        for cid in chapter_ids
    ]
    missing = {
        chapter.chapter_id: chapter.missing
        for chapter in chapters
        if chapter.missing
    }
    if missing and not allow_partial:
        return ExportResult(chapters=[], missing_by_chapter=missing, partial=False)
    return ExportResult(
        chapters=[chapter for chapter in chapters if chapter.scenes],
        missing_by_chapter=missing,
        partial=bool(missing),
    )


def document_title(meta: dict[str, Any], scope: str, chapter: ChapterExport | None) -> str:
    title = str(meta.get("title") or "Export")
    if scope == "chapter" and chapter is not None:
        return f"{title} - Kapitel {chapter.chapter_id}"
    return title


def make_cover(
    work_dir: Path,
    title: str,
    author: str,
    style: str,
    scope_label: str,
    meta: dict[str, Any],
) -> Path:
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError as exc:
        raise RuntimeError(
            "Pillow ist fuer Platzhalter-Cover nicht installiert. "
            "Bitte `pip install -r requirements.txt` ausfuehren."
        ) from exc

    cover_cfg = meta.get("cover", {}) or {}
    bg = cover_cfg.get("background", "#f59e0b")
    fg = cover_cfg.get("foreground", "#ffffff")
    img = Image.new("RGB", (1600, 2400), bg)
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 92)
        author_font = ImageFont.truetype("arial.ttf", 54)
        small_font = ImageFont.truetype("arial.ttf", 38)
    except OSError:
        title_font = ImageFont.load_default()
        author_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    def wrap(text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
        words = text.split()
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            bbox = draw.textbbox((0, 0), candidate, font=font)
            if bbox[2] - bbox[0] <= width or not current:
                current = candidate
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines

    y = 520
    for line in wrap(title, title_font, 1240):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((1600 - (bbox[2] - bbox[0])) / 2, y), line, fill=fg, font=title_font)
        y += 112
    y += 90
    for line in wrap(author, author_font, 1200):
        bbox = draw.textbbox((0, 0), line, font=author_font)
        draw.text(((1600 - (bbox[2] - bbox[0])) / 2, y), line, fill=fg, font=author_font)
        y += 70
    footer = f"{scope_label} | {style}"
    bbox = draw.textbbox((0, 0), footer, font=small_font)
    draw.text(((1600 - (bbox[2] - bbox[0])) / 2, 2040), footer, fill=fg, font=small_font)
    path = work_dir / "cover.png"
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return path


def resolve_cover_image(path_text: str, base_dir: Path | None = None) -> Path:
    raw = Path(path_text)
    path = raw if raw.is_absolute() else (base_dir or REPO_ROOT) / raw
    if not path.exists():
        raise FileNotFoundError(f"Coverbild nicht gefunden: {path}")
    if not path.is_file():
        raise FileNotFoundError(f"Coverpfad ist keine Datei: {path}")
    return path


def prepare_cover(
    work_dir: Path,
    title: str,
    author: str,
    style: str,
    scope_label: str,
    meta: dict[str, Any],
) -> Path:
    cover_cfg = meta.get("cover", {}) or {}
    image_path = str(cover_cfg.get("image_path") or "").strip()
    mode = str(cover_cfg.get("mode") or "placeholder").strip().lower()
    if mode == "image" or image_path:
        base_dir = Path(str(meta.get("_base_dir") or REPO_ROOT))
        return resolve_cover_image(image_path, base_dir)
    return make_cover(work_dir, title, author, style, scope_label, meta)


def markdown_image_path(path: Path, markdown_path: Path) -> str:
    try:
        rel = path.resolve().relative_to(markdown_path.parent.resolve())
        text = rel.as_posix()
    except ValueError:
        text = path.resolve().as_posix()
    return text.replace(" ", "%20")


def front_matter_config(meta: dict[str, Any]) -> dict[str, Any]:
    return meta.get("front_matter", {}) or {}


def should_show(meta: dict[str, Any], key: str, default: bool = True) -> bool:
    return bool(front_matter_config(meta).get(key, default))


def display_config(meta: dict[str, Any]) -> dict[str, Any]:
    return meta.get("display", {}) or {}


IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".webp")


def illustrations_config(meta: dict[str, Any]) -> dict[str, Any]:
    return meta.get("illustrations", {}) or {}


def illustrations_enabled(meta: dict[str, Any]) -> bool:
    cfg = illustrations_config(meta)
    return bool(cfg.get("enabled", False))


def book_base_dir(meta: dict[str, Any]) -> Path:
    return Path(str(meta.get("_base_dir") or REPO_ROOT))


def find_named_image(directory: Path, stem: str) -> Path | None:
    for ext in IMAGE_EXTENSIONS:
        candidate = directory / f"{stem}{ext}"
        if candidate.is_file():
            return candidate
    return None


def chapter_illustration(meta: dict[str, Any], chapter_id: str) -> Illustration | None:
    cfg = illustrations_config(meta)
    if not illustrations_enabled(meta) or not bool(cfg.get("chapter_images", True)):
        return None
    path = find_named_image(
        book_base_dir(meta) / "assets" / "chapter",
        f"chapter-{chapter_id}",
    )
    if path is None:
        return None
    return Illustration("chapter", chapter_id, None, path)


def scene_illustration(
    meta: dict[str, Any],
    chapter_id: str,
    scene_number: int,
) -> Illustration | None:
    cfg = illustrations_config(meta)
    if not illustrations_enabled(meta) or not bool(cfg.get("scene_images", True)):
        return None
    path = find_named_image(
        book_base_dir(meta) / "assets" / "scene" / chapter_id,
        f"scene-{scene_number:03d}",
    )
    if path is None:
        return None
    return Illustration("scene", chapter_id, scene_number, path)


def collect_illustrations(
    chapters: list[ChapterExport],
    meta: dict[str, Any],
) -> list[Illustration]:
    if not illustrations_enabled(meta):
        return []
    found: list[Illustration] = []
    for chapter in chapters:
        if illustration := chapter_illustration(meta, chapter.chapter_id):
            found.append(illustration)
        for scene in chapter.scenes:
            if illustration := scene_illustration(meta, chapter.chapter_id, scene.number):
                found.append(illustration)
    return found


def illustration_key(illustration: Illustration) -> str:
    if illustration.scene_number is None:
        return f"chapter:{illustration.chapter_id}"
    return f"scene:{illustration.chapter_id}:{illustration.scene_number:02d}"


def illustration_map(illustrations: list[Illustration]) -> dict[str, Illustration]:
    return {illustration_key(item): item for item in illustrations}


def illustration_label(illustration: Illustration) -> str:
    if illustration.scene_number is None:
        return f"Kapitel {illustration.chapter_id}"
    return f"Kapitel {illustration.chapter_id}, Szene {illustration.scene_number:02d}"


def render_markdown_illustration(
    illustration: Illustration,
    markdown_path: Path,
    css_classes: str,
) -> list[str]:
    ref = markdown_image_path(illustration.path, markdown_path)
    alt = illustration_label(illustration)
    return [
        f'<div class="{html.escape(css_classes)}">',
        f'  <img src="{html.escape(ref)}" alt="{html.escape(alt)}" />',
        "</div>",
        "",
    ]


def cardinal_de_ascii(value: int) -> str:
    units = {
        1: "ein",
        2: "zwei",
        3: "drei",
        4: "vier",
        5: "fuenf",
        6: "sechs",
        7: "sieben",
        8: "acht",
        9: "neun",
        10: "zehn",
        11: "elf",
        12: "zwoelf",
        13: "dreizehn",
        14: "vierzehn",
        15: "fuenfzehn",
        16: "sechzehn",
        17: "siebzehn",
        18: "achtzehn",
        19: "neunzehn",
    }
    tens = {
        20: "zwanzig",
        30: "dreissig",
        40: "vierzig",
        50: "fuenfzig",
        60: "sechzig",
        70: "siebzig",
        80: "achtzig",
        90: "neunzig",
    }
    if value <= 0:
        return str(value)
    if value < 20:
        return units[value]
    if value < 100:
        ten = value // 10 * 10
        unit = value % 10
        if unit == 0:
            return tens[ten]
        return f"{units[unit]}und{tens[ten]}"
    if value < 1000:
        hundred = value // 100
        rest = value % 100
        prefix = "hundert" if hundred == 1 else f"{units.get(hundred, str(hundred))}hundert"
        return prefix if rest == 0 else f"{prefix}{cardinal_de_ascii(rest)}"
    return str(value)


def ordinal_de_ascii(value: int) -> str:
    irregular = {
        1: "erstes",
        2: "zweites",
        3: "drittes",
        4: "viertes",
        5: "fuenftes",
        6: "sechstes",
        7: "siebtes",
        8: "achtes",
        9: "neuntes",
        10: "zehntes",
        11: "elftes",
        12: "zwoelftes",
        13: "dreizehntes",
        14: "vierzehntes",
        15: "fuenfzehntes",
        16: "sechzehntes",
        17: "siebzehntes",
        18: "achtzehntes",
        19: "neunzehntes",
    }
    text = irregular.get(value)
    if text is None:
        text = f"{cardinal_de_ascii(value)}stes"
    return text[:1].upper() + text[1:]


def display_chapter_title(chapter: ChapterExport, meta: dict[str, Any] | None = None) -> str:
    if not meta:
        return clean_chapter_title(chapter)
    chapter_cfg = (display_config(meta).get("chapters") or {})
    explicit_titles = chapter_cfg.get("titles") or {}
    if isinstance(explicit_titles, dict):
        explicit = explicit_titles.get(str(chapter.chapter_id))
        if explicit:
            return str(explicit)
    fmt = str(chapter_cfg.get("format") or "").strip()
    if not fmt:
        return clean_chapter_title(chapter)
    try:
        number = int(chapter.chapter_id)
    except ValueError:
        number = 0
    suffix = str(chapter_cfg.get("suffix") or "")
    if fmt == "words_de":
        title = f"{ordinal_de_ascii(number)}{suffix}"
    elif fmt == "number_dot":
        title = f"{number}."
    elif fmt == "number":
        title = str(number)
    else:
        title = clean_chapter_title(chapter)
    if chapter_cfg.get("include_source_title"):
        source_title = clean_chapter_title(chapter)
        if source_title and source_title != f"Kapitel {chapter.chapter_id}":
            title = f"{title}: {source_title}"
    return title


def clean_chapter_title(chapter: ChapterExport) -> str:
    title = chapter.title.strip()
    title = re.sub(r"^Kapitel\s+\d+\s*:\s*", "", title, flags=re.IGNORECASE)
    has_cyrillic = bool(re.search(r"[\u0400-\u04ff]", title))
    looks_mojibake = any(token in title for token in ("\u00d0", "\u00d1", "\u00c3"))
    if not title or has_cyrillic or looks_mojibake:
        return f"Kapitel {chapter.chapter_id}"
    return f"Kapitel {chapter.chapter_id}: {title}"


def chapter_heading_markdown(chapter: ChapterExport) -> str:
    return f"# {clean_chapter_title(chapter)} {{#kapitel-{chapter.chapter_id}}}"


def chapter_heading_markdown_for_level(
    chapter: ChapterExport,
    level: int,
    meta: dict[str, Any] | None = None,
) -> str:
    level = max(1, min(level, 6))
    chapter_cfg = (display_config(meta or {}).get("chapters") or {})
    classes = []
    if chapter_cfg:
        classes.append("chapter-heading")
        if chapter_cfg.get("align") == "center":
            classes.append("centered")
    attr_bits = [f"#kapitel-{chapter.chapter_id}", *[f".{cls}" for cls in classes]]
    attrs = " ".join(attr_bits)
    return f"{'#' * level} {display_chapter_title(chapter, meta)} {{{attrs}}}"


def group_for_chapter(meta: dict[str, Any], chapter_id: str) -> dict[str, Any] | None:
    groups = meta.get("structure_groups") or []
    for group in groups:
        start = str(group.get("from") or "")
        end = str(group.get("to") or "")
        if start and end and start <= chapter_id <= end:
            return group
    return None


def group_heading_markdown(group: dict[str, Any]) -> str:
    label = str(group.get("label") or group.get("id") or "Buch").strip()
    group_id = sanitize_filename(str(group.get("id") or label))
    return f"# {label} {{#gruppe-{group_id}}}"


def text_blocks(value: Any) -> list[str]:
    if not value:
        return []
    if isinstance(value, str):
        return [block.strip() for block in re.split(r"\n\s*\n", value) if block.strip()]
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [str(value).strip()]


def translator_label(meta: dict[str, Any]) -> str:
    return str(meta.get("translator_label") or "Uebersetzung")


def add_imprint_lines(
    lines: list[str],
    title: str,
    meta: dict[str, Any],
    scope: str,
) -> None:
    custom_text = str(meta.get("imprint_text") or "").strip()
    if custom_text:
        lines.extend([custom_text, ""])
        return

    lines.append(f"**Titel:** {title}")
    lines.append("")
    if meta.get("subtitle") and scope == "book":
        lines.extend([f"**Untertitel:** {meta['subtitle']}", ""])
    if meta.get("author"):
        lines.extend([f"**Autor:** {meta['author']}", ""])
    if meta.get("translator"):
        lines.extend([f"**{translator_label(meta)}:** {meta['translator']}", ""])
    if meta.get("publisher"):
        lines.extend([f"**Herausgeber:** {meta['publisher']}", ""])
    if meta.get("rights"):
        lines.extend([f"**Rechte:** {meta['rights']}", ""])
    lines.extend([f"**Sprache:** {meta.get('language', 'de-DE')}", ""])


def html_paragraphs(text: str) -> str:
    paragraphs = []
    for paragraph in re.split(r"\n\s*\n", text.strip()):
        clean = " ".join(line.strip() for line in paragraph.splitlines()).strip()
        if clean:
            paragraphs.append(f"<p>{html.escape(clean)}</p>")
    return "\n".join(paragraphs)


def html_imprint_paragraph(block: str, css_class: str) -> str:
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if not lines:
        return ""
    body = "<br />\n    ".join(html.escape(line) for line in lines)
    return f'  <p class="{css_class}">{body}</p>'


def imprint_block_class(index: int, block: str) -> str:
    if index == 0:
        return "imprint-title"
    if index == 1:
        return "imprint-author"
    if index == 2:
        return "imprint-subtitle"
    lowered = block.casefold()
    rights_markers = ("©", "rechte", "gemeinfrei", "vorbehalten")
    if any(marker in lowered for marker in rights_markers):
        return "imprint-rights"
    return "imprint-meta"


def render_imprint_page_html(
    heading: str,
    title: str,
    meta: dict[str, Any],
    scope: str,
) -> list[str]:
    blocks = text_blocks(meta.get("imprint_text"))
    rows: list[tuple[str, str]]
    if blocks:
        rows = [
            (block, imprint_block_class(index, block))
            for index, block in enumerate(blocks)
        ]
    else:
        fallback_blocks = [
            title,
            str(meta.get("subtitle", "")) if scope == "book" else "",
            str(meta.get("author", "")),
            f"{translator_label(meta)}: {meta['translator']}" if meta.get("translator") else "",
            f"Herausgeber: {meta['publisher']}" if meta.get("publisher") else "",
            f"Rechte: {meta['rights']}" if meta.get("rights") else "",
            f"Sprache: {meta.get('language', 'de-DE')}",
        ]
        rows = [(block, "imprint-meta") for block in fallback_blocks if block]
    lines = [
        f'<section id="frontmatter-imprint" class="frontmatter-page imprintpage" epub:type="copyright-page">',
        f"  <h1>{html.escape(heading)}</h1>",
    ]
    for block, css_class in rows:
        paragraph = html_imprint_paragraph(block, css_class)
        if paragraph:
            lines.append(paragraph)
    lines.extend(["</section>", ""])
    return lines


def render_title_page_html(
    title: str,
    meta: dict[str, Any],
    scope: str,
    partial: bool,
) -> list[str]:
    lines = [
        '::: {.frontmatter-page .titlepage epub:type="titlepage"}',
        f"[{title}]{{.book-title}}",
        "",
    ]
    if meta.get("subtitle") and scope == "book":
        lines.extend([f"[{meta['subtitle']}]{{.subtitle}}", ""])
    if meta.get("author"):
        lines.extend([f"[{meta['author']}]{{.author}}", ""])
    for block in text_blocks(meta.get("title_page_extra")):
        lines.extend([f"[{block}]{{.title-extra}}", ""])
    if meta.get("translator"):
        lines.append(
            f"[{translator_label(meta)}: {meta['translator']}]{{.translator}}"
        )
        lines.append("")
    if partial:
        lines.extend(["[Teil-Export: Es fehlen noch Szenen.]{.partial-note}", ""])
    lines.extend([":::", ""])
    return lines


def render_frontmatter_page_html(
    heading: str,
    body: str,
    anchor: str,
    epub_type: str,
) -> list[str]:
    return [
        f'<section id="{html.escape(anchor)}" class="frontmatter-page textpage" epub:type="{html.escape(epub_type)}">',
        f"  <h1>{html.escape(heading)}</h1>",
        html_paragraphs(body),
        "</section>",
        "",
    ]


def render_front_matter_markdown(
    meta: dict[str, Any],
    title: str,
    scope: str,
    partial: bool,
    cover_ref: str | None,
) -> list[str]:
    fm = front_matter_config(meta)
    description = str(meta.get("description") or "").strip()
    explicit_summary = str(meta.get("summary") or "").strip()
    summary = explicit_summary
    author_bio = str(meta.get("author_bio") or "").strip()
    # EPUB uses Pandoc's official cover image. Do not also emit a Markdown
    # cover chapter, otherwise readers show duplicate cover/title fragments.
    wants_cover = False
    wants_title = should_show(meta, "title_page", True)
    wants_summary = bool(summary and should_show(meta, "summary_page", True))
    wants_author_bio = bool(author_bio and should_show(meta, "author_bio_page", True))
    wants_description = bool(
        description
        and not explicit_summary
        and should_show(meta, "description_page", True)
    )
    wants_imprint = should_show(meta, "imprint_page", True)
    lines: list[str] = []
    if not (wants_cover or wants_title or wants_summary or wants_author_bio or wants_description or wants_imprint or partial):
        return lines

    if fm.get("combined_epub_front_matter", True):
        lines.extend([f"# {fm.get('combined_heading', 'Titelei')} {{#frontmatter}}", ""])
        if wants_cover:
            lines.extend([f"![Cover]({cover_ref})", ""])
        if wants_title:
            lines.extend([f"## {fm.get('title_heading', 'Titelseite')}", ""])
            lines.extend([f"**{title}**", ""])
            if meta.get("subtitle") and scope == "book":
                lines.extend([str(meta["subtitle"]), ""])
            if meta.get("author"):
                lines.extend([str(meta["author"]), ""])
            for block in text_blocks(meta.get("title_page_extra")):
                lines.extend([block, ""])
            if meta.get("translator"):
                lines.extend([f"{translator_label(meta)}: {meta['translator']}", ""])
        if partial:
            lines.extend(["> Teil-Export: Es fehlen noch Szenen.", ""])
        if wants_summary:
            lines.extend([f"## {fm.get('summary_heading', fm.get('description_heading', 'Zusammenfassung'))}", "", summary, ""])
        if wants_author_bio:
            lines.extend([f"## {fm.get('author_bio_heading', 'Leben des Autors')}", "", author_bio, ""])
        if wants_description:
            lines.extend([f"## {fm.get('description_heading', 'Zu dieser Ausgabe')}", "", description, ""])
        if wants_imprint:
            lines.extend([f"## {fm.get('imprint_heading', 'Impressum')}", ""])
            add_imprint_lines(lines, title, meta, scope)
        return lines

    if wants_title:
        lines.extend(render_title_page_html(title, meta, scope, partial))
    if partial:
        # The title page already carries the partial note when present.
        if not wants_title:
            lines.extend(["> Teil-Export: Es fehlen noch Szenen.", ""])
    if wants_summary:
        heading = fm.get("summary_heading", fm.get("description_heading", "Zusammenfassung"))
        lines.extend(render_frontmatter_page_html(str(heading), summary, "frontmatter-summary", "preface"))
    if wants_author_bio:
        heading = fm.get("author_bio_heading", "Leben des Autors")
        lines.extend(render_frontmatter_page_html(str(heading), author_bio, "frontmatter-author", "foreword"))
    if wants_description:
        heading = fm.get("description_heading", "Zu dieser Ausgabe")
        lines.extend([f"# {heading} {{#frontmatter-description}}", "", description, ""])
    if wants_imprint:
        heading = fm.get("imprint_heading", "Impressum")
        lines.extend(render_imprint_page_html(str(heading), title, meta, scope))
    return lines


def render_export_markdown(
    chapters: list[ChapterExport],
    meta: dict[str, Any],
    style: str,
    scope: str,
    partial: bool,
    cover_ref: str | None,
    markdown_path: Path | None = None,
    illustrations: list[Illustration] | None = None,
) -> str:
    title = document_title(meta, scope, chapters[0] if chapters else None)
    illustrations_by_key = illustration_map(illustrations or [])
    lines = [
        "---",
        f"title: {yaml_scalar(title)}",
        f"author: {yaml_scalar(meta.get('author', ''))}",
        f"lang: {meta.get('language', 'de-DE')}",
        f"date: {date.today().isoformat()}",
    ]
    if meta.get("rights"):
        lines.append(f"rights: {yaml_scalar(meta['rights'])}")
    lines.extend(["---", ""])
    lines.extend(render_front_matter_markdown(meta, title, scope, partial, cover_ref))
    illustration_cfg = illustrations_config(meta)
    if should_show(meta, "toc_page", False):
        fm = front_matter_config(meta)
        heading = fm.get("toc_heading", "Inhalt")
        lines.extend([f"# {heading} {{#frontmatter-toc}}", ""])
        for chapter in chapters:
            lines.append(f"- [{display_chapter_title(chapter, meta)}](#kapitel-{chapter.chapter_id})")
        lines.append("")
    last_group_id = None
    has_groups = bool(meta.get("structure_groups"))
    chapter_level = 2 if has_groups else 1
    for chapter in chapters:
        group = group_for_chapter(meta, chapter.chapter_id)
        group_id = group.get("id") if group else None
        if group and group_id != last_group_id:
            lines.append(group_heading_markdown(group))
            lines.append("")
            last_group_id = group_id
        chapter_image = illustrations_by_key.get(f"chapter:{chapter.chapter_id}")
        if chapter_image and markdown_path is not None:
            lines.extend(
                render_markdown_illustration(
                    chapter_image,
                    markdown_path,
                    (
                        "chapter-illustration page-break-after"
                        if illustration_cfg.get("chapter_page_break_after_image", True)
                        else "chapter-illustration"
                    ),
                )
            )
        lines.append(chapter_heading_markdown_for_level(chapter, chapter_level, meta))
        lines.append("")
        scene_cfg = (display_config(meta).get("scenes") or {})
        show_scene_marker = bool(scene_cfg.get("show"))
        has_display = bool(display_config(meta))
        for idx, scene in enumerate(chapter.scenes):
            scene_image = illustrations_by_key.get(
                f"scene:{chapter.chapter_id}:{scene.number:02d}"
            )
            if scene_image and markdown_path is not None:
                lines.extend(
                    render_markdown_illustration(
                        scene_image,
                        markdown_path,
                        (
                            "scene-illustration page-break-after"
                            if illustration_cfg.get("scene_page_break_after_image", False)
                            else "scene-illustration"
                        ),
                    )
                )
            if show_scene_marker:
                marker = str(scene.number if scene_cfg.get("format", "number") == "number" else scene.number)
                align_class = " .centered" if scene_cfg.get("align") == "center" else ""
                lines.extend([f"[{marker}]{{.scene-marker{align_class}}}", ""])
            elif idx and not has_display:
                lines.extend(["", str(meta.get("output", {}).get("scene_separator", "* * *")), ""])
            elif idx and str(scene_cfg.get("separator") or ""):
                lines.extend(["", str(scene_cfg.get("separator")), ""])
            lines.append(scene.text)
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def html_asset_path(path: Path, html_path: Path) -> str:
    return markdown_image_path(path, html_path)


def inline_markdown_to_html(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", escaped)
    return escaped


def markdown_fragment_to_html(text: str) -> str:
    try:
        import markdown as markdown_lib
    except ImportError:
        blocks: list[str] = []
        for raw_block in re.split(r"\n\s*\n", text.strip()):
            block = raw_block.strip()
            if not block:
                continue
            quote_lines = [
                line.strip()[1:].strip()
                for line in block.splitlines()
                if line.strip().startswith(">")
            ]
            if quote_lines and len(quote_lines) == len(block.splitlines()):
                body = " ".join(quote_lines)
                blocks.append(f"<blockquote><p>{inline_markdown_to_html(body)}</p></blockquote>")
                continue
            heading = re.match(r"^(#{1,6})\s+(.+)$", block)
            if heading:
                level = min(len(heading.group(1)), 3)
                blocks.append(f"<h{level + 1}>{inline_markdown_to_html(heading.group(2).strip())}</h{level + 1}>")
            else:
                body = " ".join(line.strip() for line in block.splitlines())
                blocks.append(f"<p>{inline_markdown_to_html(body)}</p>")
        return "\n".join(blocks)
    return markdown_lib.markdown(
        text,
        extensions=["extra", "sane_lists"],
        output_format="html5",
    )


def render_pdf_image(path: Path, html_path: Path, css_class: str, alt: str = "") -> str:
    ref = html_asset_path(path, html_path)
    return (
        f'<div class="{html.escape(css_class)}">'
        f'<img src="{html.escape(ref)}" alt="{html.escape(alt)}" />'
        "</div>"
    )


def render_pdf_front_matter(
    meta: dict[str, Any],
    title: str,
    scope: str,
    partial: bool,
    cover_path: Path,
    html_path: Path,
) -> list[str]:
    fm = front_matter_config(meta)
    description = str(meta.get("description") or "").strip()
    explicit_summary = str(meta.get("summary") or "").strip()
    summary = explicit_summary
    author_bio = str(meta.get("author_bio") or "").strip()
    lines: list[str] = []

    if should_show(meta, "cover_in_body", True):
        lines.extend([
            '<section class="frontmatter-page coverpage" aria-label="Cover">',
            render_pdf_image(cover_path, html_path, "cover-image", "Cover"),
            "</section>",
        ])

    if should_show(meta, "title_page", True):
        lines.extend([
            '<section class="frontmatter-page titlepage">',
            f'<h1 class="book-title">{html.escape(title)}</h1>',
        ])
        if meta.get("subtitle") and scope == "book":
            lines.append(f'<p class="subtitle">{html.escape(str(meta["subtitle"]))}</p>')
        if meta.get("author"):
            lines.append(f'<p class="author">{html.escape(str(meta["author"]))}</p>')
        for block in text_blocks(meta.get("title_page_extra")):
            lines.append(f'<p class="title-extra">{html.escape(block)}</p>')
        if meta.get("translator"):
            label = translator_label(meta)
            lines.append(
                f'<p class="translator">{html.escape(label)}: '
                f'{html.escape(str(meta["translator"]))}</p>'
            )
        if partial:
            lines.append('<p class="partial-note">Teil-Export: Es fehlen noch Szenen.</p>')
        lines.append("</section>")

    if summary and should_show(meta, "summary_page", True):
        heading = fm.get("summary_heading", fm.get("description_heading", "Zusammenfassung"))
        lines.extend([
            '<section class="frontmatter-page textpage" id="frontmatter-summary">',
            f"<h1>{html.escape(str(heading))}</h1>",
            html_paragraphs(summary),
            "</section>",
        ])

    if author_bio and should_show(meta, "author_bio_page", True):
        heading = fm.get("author_bio_heading", "Leben des Autors")
        lines.extend([
            '<section class="frontmatter-page textpage" id="frontmatter-author">',
            f"<h1>{html.escape(str(heading))}</h1>",
            html_paragraphs(author_bio),
            "</section>",
        ])

    wants_description = (
        description
        and not explicit_summary
        and should_show(meta, "description_page", True)
    )
    if wants_description:
        heading = fm.get("description_heading", "Zu dieser Ausgabe")
        lines.extend([
            '<section class="frontmatter-page textpage" id="frontmatter-description">',
            f"<h1>{html.escape(str(heading))}</h1>",
            html_paragraphs(description),
            "</section>",
        ])

    if should_show(meta, "imprint_page", True):
        heading = fm.get("imprint_heading", "Impressum")
        rows = render_imprint_page_html(str(heading), title, meta, scope)
        html_rows = [
            line.replace("frontmatter-page imprintpage", "frontmatter-page imprintpage")
            for line in rows
        ]
        lines.extend(html_rows)

    return lines


def render_pdf_html(
    chapters: list[ChapterExport],
    meta: dict[str, Any],
    style: str,
    scope: str,
    partial: bool,
    cover_path: Path,
    html_path: Path,
    css_path: Path,
    illustrations: list[Illustration] | None = None,
) -> str:
    title = document_title(meta, scope, chapters[0] if chapters else None)
    illustrations_by_key = illustration_map(illustrations or [])
    css_ref = html_asset_path(css_path, html_path)
    lang = html.escape(str(meta.get("language", "de-DE")))
    lines = [
        "<!doctype html>",
        f'<html lang="{lang}">',
        "<head>",
        '  <meta charset="utf-8" />',
        f"  <title>{html.escape(title)}</title>",
        f'  <link rel="stylesheet" href="{html.escape(css_ref)}" />',
        "</head>",
        "<body>",
        '<main class="book">',
    ]
    lines.extend(render_pdf_front_matter(meta, title, scope, partial, cover_path, html_path))

    if should_show(meta, "toc_page", False):
        fm = front_matter_config(meta)
        heading = fm.get("toc_heading", "Inhalt")
        lines.extend([
            '<section class="frontmatter-page tocpage" id="frontmatter-toc">',
            f"<h1>{html.escape(str(heading))}</h1>",
            "<ol>",
        ])
        for chapter in chapters:
            lines.append(
                f'<li><a href="#kapitel-{html.escape(chapter.chapter_id)}">'
                f"{html.escape(display_chapter_title(chapter, meta))}</a></li>"
            )
        lines.extend(["</ol>", "</section>"])

    last_group_id = None
    has_groups = bool(meta.get("structure_groups"))
    scene_cfg = (display_config(meta).get("scenes") or {})
    illustration_cfg = illustrations_config(meta)
    show_scene_marker = bool(scene_cfg.get("show"))
    has_display = bool(display_config(meta))
    for chapter in chapters:
        group = group_for_chapter(meta, chapter.chapter_id)
        group_id = group.get("id") if group else None
        if group and group_id != last_group_id:
            label = str(group.get("label") or group_id)
            lines.append(
                f'<section class="group-heading" id="gruppe-{sanitize_filename(str(group_id or label))}">'
                f"<h1>{html.escape(label)}</h1></section>"
            )
            last_group_id = group_id

        lines.append(f'<section class="chapter" id="kapitel-{html.escape(chapter.chapter_id)}">')
        chapter_image = illustrations_by_key.get(f"chapter:{chapter.chapter_id}")
        if chapter_image:
            css_class = "chapter-illustration"
            if illustration_cfg.get("chapter_page_break_after_image", True):
                css_class += " page-break-after"
            lines.append(
                render_pdf_image(
                    chapter_image.path,
                    html_path,
                    css_class,
                    illustration_label(chapter_image),
                )
            )
        level = "h2" if has_groups else "h1"
        lines.append(
            f'<{level} class="chapter-heading">'
            f"{html.escape(display_chapter_title(chapter, meta))}</{level}>"
        )
        for idx, scene in enumerate(chapter.scenes):
            scene_image = illustrations_by_key.get(
                f"scene:{chapter.chapter_id}:{scene.number:02d}"
            )
            if scene_image:
                css_class = "scene-illustration"
                if illustration_cfg.get("scene_page_break_after_image", False):
                    css_class += " page-break-after"
                lines.append(
                    render_pdf_image(
                        scene_image.path,
                        html_path,
                        css_class,
                        illustration_label(scene_image),
                    )
                )
            if show_scene_marker:
                marker = str(scene.number if scene_cfg.get("format", "number") == "number" else scene.number)
                lines.append(f'<div class="scene-marker">{html.escape(marker)}</div>')
            elif idx and not has_display:
                sep_text = str(meta.get("output", {}).get("scene_separator", "* * *"))
                lines.append(f'<div class="scene-separator">{html.escape(sep_text)}</div>')
            elif idx and str(scene_cfg.get("separator") or ""):
                lines.append(f'<div class="scene-separator">{html.escape(str(scene_cfg.get("separator")))}</div>')
            lines.append(f'<div class="scene-text">{markdown_fragment_to_html(scene.text)}</div>')
        lines.append("</section>")

    lines.extend(["</main>", "</body>", "</html>"])
    return "\n".join(lines) + "\n"


def write_pdf_css(path: Path) -> Path:
    css = (
        "@page {\n"
        "  size: A5;\n"
        "  margin: 18mm 16mm 20mm 16mm;\n"
        "}\n"
        "\n"
        "* { box-sizing: border-box; }\n"
        "\n"
        "html, body {\n"
        "  margin: 0;\n"
        "  padding: 0;\n"
        "}\n"
        "\n"
        "body {\n"
        '  font-family: Georgia, "Times New Roman", serif;\n'
        "  font-size: 10.5pt;\n"
        "  line-height: 1.45;\n"
        "  color: #202027;\n"
        "  -webkit-print-color-adjust: exact;\n"
        "  print-color-adjust: exact;\n"
        "}\n"
        "\n"
        "a { color: inherit; text-decoration: none; }\n"
        "\n"
        ".frontmatter-page,\n"
        ".chapter,\n"
        ".group-heading {\n"
        "  break-before: page;\n"
        "  page-break-before: always;\n"
        "}\n"
        "\n"
        ".coverpage,\n"
        ".titlepage {\n"
        "  min-height: 210mm;\n"
        "  display: flex;\n"
        "  flex-direction: column;\n"
        "  align-items: center;\n"
        "  justify-content: center;\n"
        "  text-align: center;\n"
        "}\n"
        "\n"
        ".cover-image img {\n"
        "  max-width: 100%;\n"
        "  max-height: 175mm;\n"
        "  object-fit: contain;\n"
        "}\n"
        "\n"
        ".book-title {\n"
        "  margin: 0 0 0.7em;\n"
        "  font-size: 25pt;\n"
        "  font-weight: normal;\n"
        "  line-height: 1.12;\n"
        "}\n"
        "\n"
        ".subtitle { margin: 0.3em 0 2em; font-size: 12pt; }\n"
        ".author { margin: 1.8em 0 0.5em; font-size: 13pt; }\n"
        ".translator, .title-extra, .partial-note { margin: 0.55em 0 0; font-size: 9.5pt; }\n"
        "\n"
        ".textpage, .imprintpage, .tocpage {\n"
        "  max-width: 100%;\n"
        "}\n"
        "\n"
        ".textpage h1, .imprintpage h1, .tocpage h1 {\n"
        "  margin: 0 0 12mm;\n"
        "  font-size: 15pt;\n"
        "  font-weight: normal;\n"
        "  text-align: center;\n"
        "}\n"
        "\n"
        ".imprintpage { padding-top: 12mm; }\n"
        ".imprintpage p { margin: 0 0 0.8em; text-indent: 0; }\n"
        ".imprint-title { font-size: 16pt; }\n"
        ".imprint-author { font-size: 12pt; margin-bottom: 1.5em; }\n"
        ".imprint-subtitle { font-size: 10.5pt; margin-bottom: 1.8em; }\n"
        ".imprint-meta { font-size: 9.2pt; }\n"
        ".imprint-rights { font-size: 8.4pt; color: #444; }\n"
        "\n"
        ".tocpage ol { margin: 0; padding-left: 1.4em; }\n"
        ".tocpage li { margin: 0 0 0.45em; }\n"
        "\n"
        ".group-heading h1 {\n"
        "  margin-top: 58mm;\n"
        "  text-align: center;\n"
        "  font-size: 17pt;\n"
        "  font-weight: normal;\n"
        "}\n"
        "\n"
        ".chapter-heading {\n"
        "  margin: 30mm 0 13mm;\n"
        "  text-align: center;\n"
        "  font-size: 15pt;\n"
        "  font-weight: normal;\n"
        "  line-height: 1.25;\n"
        "}\n"
        "\n"
        ".chapter-illustration,\n"
        ".scene-illustration {\n"
        "  margin: 0 0 11mm;\n"
        "  text-align: center;\n"
        "  break-inside: avoid;\n"
        "  page-break-inside: avoid;\n"
        "}\n"
        "\n"
        ".chapter-illustration img,\n"
        ".scene-illustration img {\n"
        "  max-width: 100%;\n"
        "  max-height: 150mm;\n"
        "  height: auto;\n"
        "}\n"
        "\n"
        ".page-break-after {\n"
        "  break-after: page;\n"
        "  page-break-after: always;\n"
        "}\n"
        "\n"
        ".scene-marker {\n"
        "  margin: 1.15em 0 0.95em;\n"
        "  text-align: center;\n"
        "  font-size: 10pt;\n"
        "}\n"
        "\n"
        ".scene-separator {\n"
        "  margin: 1.6em 0;\n"
        "  text-align: center;\n"
        "  letter-spacing: 0.08em;\n"
        "}\n"
        "\n"
        ".scene-text p,\n"
        ".textpage p {\n"
        "  margin: 0 0 0.85em;\n"
        "  orphans: 2;\n"
        "  widows: 2;\n"
        "}\n"
        "\n"
        ".scene-text blockquote {\n"
        "  margin: 0 0 1.1em;\n"
        "  font-size: 9.5pt;\n"
        "  font-style: italic;\n"
        "  color: #444;\n"
        "}\n"
        "\n"
        ".scene-text h2,\n"
        ".scene-text h3,\n"
        ".scene-text h4 {\n"
        "  margin: 1.25em 0 0.7em;\n"
        "  font-size: 11.5pt;\n"
        "  font-weight: normal;\n"
        "}\n"
        "\n"
        "figure, img { break-inside: avoid; page-break-inside: avoid; }\n"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(css + "\n", encoding="utf-8")
    return path


def write_pdf(path: Path, html_path: Path) -> None:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise RuntimeError(
            "Playwright ist fuer den PDF-Export nicht installiert. "
            "Bitte `pip install -r requirements.txt` und danach "
            "`python -m playwright install chromium` ausfuehren."
        ) from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch()
        try:
            page = browser.new_page()
            page.goto(html_path.resolve().as_uri(), wait_until="load")
            page.pdf(
                path=str(path),
                print_background=True,
                prefer_css_page_size=True,
                outline=True,
                tagged=True,
                margin={
                    "top": "0",
                    "right": "0",
                    "bottom": "0",
                    "left": "0",
                },
            )
        finally:
            browser.close()


def write_docx(
    path: Path,
    chapters: list[ChapterExport],
    meta: dict[str, Any],
    style: str,
    scope: str,
    partial: bool,
    cover_path: Path,
    illustrations: list[Illustration] | None = None,
) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "python-docx ist nicht installiert. "
            "Bitte `pip install -r requirements.txt` ausfuehren."
        ) from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    styles = document.styles
    styles["Normal"].font.name = "Georgia"
    styles["Normal"].font.size = Pt(11.5)
    styles["Title"].font.name = "Georgia"
    styles["Title"].font.size = Pt(28)
    styles["Heading 1"].font.name = "Georgia"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(45, 45, 60)

    title = document_title(meta, scope, chapters[0] if chapters else None)
    fm = front_matter_config(meta)
    description = str(meta.get("description") or "").strip()
    explicit_summary = str(meta.get("summary") or "").strip()
    summary = explicit_summary
    author_bio = str(meta.get("author_bio") or "").strip()
    illustrations_by_key = illustration_map(illustrations or [])
    illustration_cfg = illustrations_config(meta)

    def add_illustration_picture(
        illustration: Illustration,
        width: Any,
        page_break_after: bool,
    ) -> None:
        document.add_picture(str(illustration.path), width=width)
        picture_paragraph = document.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.space_after = Pt(10)
        if page_break_after:
            document.add_page_break()

    if should_show(meta, "cover_in_body", True):
        document.add_picture(str(cover_path), width=Inches(4.1))
        last = document.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

    if should_show(meta, "title_page", True):
        document.add_heading(title, 0)
        subtitle = meta.get("subtitle")
        if subtitle and scope == "book":
            p = document.add_paragraph(str(subtitle))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(str(meta.get("author", "")))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for block in text_blocks(meta.get("title_page_extra")):
            p = document.add_paragraph(block)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if meta.get("translator"):
            p = document.add_paragraph(f"{translator_label(meta)}: {meta['translator']}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if partial:
            p = document.add_paragraph("Teil-Export: Es fehlen noch Szenen.")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

    if summary and should_show(meta, "summary_page", True):
        document.add_heading(str(fm.get("summary_heading", fm.get("description_heading", "Zusammenfassung"))), 1)
        for paragraph in summary.split("\n\n"):
            text = paragraph.strip()
            if text:
                p = document.add_paragraph(text)
                p.paragraph_format.space_after = Pt(7)
                p.paragraph_format.line_spacing = 1.12
        document.add_page_break()

    if author_bio and should_show(meta, "author_bio_page", True):
        document.add_heading(str(fm.get("author_bio_heading", "Leben des Autors")), 1)
        for paragraph in author_bio.split("\n\n"):
            text = paragraph.strip()
            if text:
                p = document.add_paragraph(text)
                p.paragraph_format.space_after = Pt(7)
                p.paragraph_format.line_spacing = 1.12
        document.add_page_break()

    wants_description = (
        description
        and not explicit_summary
        and should_show(meta, "description_page", True)
    )
    if wants_description:
        document.add_heading(str(fm.get("description_heading", "Zu dieser Ausgabe")), 1)
        for paragraph in description.split("\n\n"):
            text = paragraph.strip()
            if text:
                p = document.add_paragraph(text)
                p.paragraph_format.space_after = Pt(7)
                p.paragraph_format.line_spacing = 1.12
        document.add_page_break()

    subtitle = meta.get("subtitle")

    if should_show(meta, "imprint_page", True):
        heading = document.add_heading(str(fm.get("imprint_heading", "Impressum")), 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_after = Pt(22)
        custom_imprint = text_blocks(meta.get("imprint_text"))
        if custom_imprint:
            for block_index, block in enumerate(custom_imprint):
                block_class = imprint_block_class(block_index, block)
                size = Pt(9.5)
                space_after = Pt(9)
                color = RGBColor(45, 45, 60)
                if block_class == "imprint-title":
                    size = Pt(18)
                    space_after = Pt(3)
                    color = RGBColor(32, 32, 39)
                elif block_class == "imprint-author":
                    size = Pt(12)
                    space_after = Pt(16)
                elif block_class == "imprint-subtitle":
                    size = Pt(10.5)
                    space_after = Pt(18)
                elif block_class == "imprint-rights":
                    size = Pt(8.8)
                    color = RGBColor(70, 70, 78)
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = space_after
                p.paragraph_format.line_spacing = 1.08
                lines = block.splitlines()
                for index, line in enumerate(lines):
                    if index:
                        p.add_run().add_break()
                    run = p.add_run(line.strip())
                    run.font.name = "Georgia"
                    run.font.size = size
                    run.font.color.rgb = color
        else:
            imprint_rows = [
                ("Titel", title),
                ("Untertitel", str(meta.get("subtitle", "")) if scope == "book" else ""),
                ("Autor", str(meta.get("author", ""))),
                (translator_label(meta), str(meta.get("translator", ""))),
                ("Herausgeber", str(meta.get("publisher", ""))),
                ("Rechte", str(meta.get("rights", ""))),
                ("Sprache", str(meta.get("language", "de-DE"))),
            ]
            for label, value in imprint_rows:
                if not value:
                    continue
                p = document.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)
        document.add_page_break()

    if should_show(meta, "toc_page", True):
        document.add_heading(str(fm.get("toc_heading", "Inhalt")), 1)
        for chapter in chapters:
            document.add_paragraph(display_chapter_title(chapter, meta), style=None)
        document.add_page_break()

    last_group_id = None
    for cidx, chapter in enumerate(chapters):
        if cidx:
            document.add_page_break()
        group = group_for_chapter(meta, chapter.chapter_id)
        group_id = group.get("id") if group else None
        if group and group_id != last_group_id:
            document.add_heading(str(group.get("label") or group_id), 1)
            last_group_id = group_id
        chapter_image = illustrations_by_key.get(f"chapter:{chapter.chapter_id}")
        if chapter_image:
            add_illustration_picture(
                chapter_image,
                Inches(5.7),
                bool(illustration_cfg.get("chapter_page_break_after_image", True)),
            )
        heading = document.add_heading(display_chapter_title(chapter, meta), 2 if group else 1)
        chapter_cfg = (display_config(meta).get("chapters") or {})
        if chapter_cfg.get("align") == "center":
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        scene_cfg = (display_config(meta).get("scenes") or {})
        show_scene_marker = bool(scene_cfg.get("show"))
        has_display = bool(display_config(meta))
        for sidx, scene in enumerate(chapter.scenes):
            scene_image = illustrations_by_key.get(
                f"scene:{chapter.chapter_id}:{scene.number:02d}"
            )
            if scene_image:
                add_illustration_picture(
                    scene_image,
                    Inches(5.2),
                    bool(illustration_cfg.get("scene_page_break_after_image", False)),
                )
            if show_scene_marker:
                marker = document.add_paragraph(str(scene.number))
                if scene_cfg.get("align") == "center":
                    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
                marker.paragraph_format.space_after = Pt(8)
            elif sidx and not has_display:
                sep_text = str(meta.get("output", {}).get("scene_separator", "* * *"))
                sep = document.add_paragraph(sep_text)
                sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif sidx and str(scene_cfg.get("separator") or ""):
                sep = document.add_paragraph(str(scene_cfg.get("separator")))
                if scene_cfg.get("align") == "center":
                    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for kind, block in markdown_to_plain_blocks(scene.text):
                if kind.startswith("heading"):
                    level = int(kind[-1]) + 1
                    document.add_heading(block, min(level, 3))
                else:
                    p = document.add_paragraph(block)
                    p.paragraph_format.first_line_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(7)
                    p.paragraph_format.line_spacing = 1.12
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def write_epub_css(path: Path) -> Path:
    css = (
        "body {\n"
        '  font-family: Georgia, "Times New Roman", serif;\n'
        "  line-height: 1.42;\n"
        "}\n"
        "\n"
        ".frontmatter-page {\n"
        "  break-before: page;\n"
        "  page-break-before: always;\n"
        "  margin: 0 auto;\n"
        "  max-width: 38em;\n"
        "}\n"
        "\n"
        ".frontmatter-page h1 {\n"
        "  margin-top: 0;\n"
        "  margin-bottom: 1em;\n"
        "  text-align: center;\n"
        "  font-size: 1.6em;\n"
        "  font-weight: normal;\n"
        "  letter-spacing: 0;\n"
        "}\n"
        "\n"
        ".frontmatter-page p {\n"
        "  text-indent: 0;\n"
        "  margin: 0 0 0.8em;\n"
        "}\n"
        "\n"
        ".frontmatter-page.textpage {\n"
        "  font-size: 0.93em;\n"
        "  line-height: 1.36;\n"
        "  max-width: 40em;\n"
        "}\n"
        "\n"
        ".frontmatter-page.textpage h1 {\n"
        "  font-size: 1.35em;\n"
        "}\n"
        "\n"
        "section.chapter-heading.centered {\n"
        "  text-align: left;\n"
        "}\n"
        "\n"
        "h1.chapter-heading,\n"
        "h2.chapter-heading,\n"
        "h3.chapter-heading {\n"
        "  font-size: 1.22em;\n"
        "  font-weight: normal;\n"
        "  margin: 1.35em 0 1em;\n"
        "}\n"
        "\n"
        "h1.chapter-heading.centered,\n"
        "h2.chapter-heading.centered,\n"
        "h3.chapter-heading.centered {\n"
        "  text-align: center;\n"
        "}\n"
        "\n"
        ".scene-marker {\n"
        "  display: block;\n"
        "  margin: 1.25em 0 1.05em;\n"
        "  text-indent: 0;\n"
        "}\n"
        "\n"
        ".scene-marker.centered {\n"
        "  text-align: center;\n"
        "}\n"
        "\n"
        ".chapter-illustration,\n"
        ".scene-illustration {\n"
        "  display: block;\n"
        "  margin: 1.2em auto;\n"
        "  text-align: center;\n"
        "}\n"
        "\n"
        ".chapter-illustration img,\n"
        ".scene-illustration img {\n"
        "  max-width: 100%;\n"
        "  height: auto;\n"
        "}\n"
        "\n"
        ".page-break-after {\n"
        "  break-after: page;\n"
        "  page-break-after: always;\n"
        "}\n"
        "\n"
        ".titlepage {\n"
        "  padding-top: 28%;\n"
        "  text-align: center;\n"
        "}\n"
        "\n"
        ".titlepage p {\n"
        "  text-align: center;\n"
        "  text-indent: 0;\n"
        "  margin-left: 0;\n"
        "  margin-right: 0;\n"
        "}\n"
        "\n"
        ".titlepage .book-title {\n"
        "  display: block;\n"
        "  margin: 0 0 0.7em;\n"
        "  font-size: 2.15em;\n"
        "  font-weight: normal;\n"
        "  line-height: 1.15;\n"
        "}\n"
        "\n"
        ".titlepage .subtitle {\n"
        "  display: block;\n"
        "  margin-top: 0.3em;\n"
        "  font-size: 1.15em;\n"
        "}\n"
        "\n"
        ".titlepage .author {\n"
        "  display: block;\n"
        "  margin-top: 2.4em;\n"
        "  font-size: 1.18em;\n"
        "}\n"
        "\n"
        ".titlepage .title-extra {\n"
        "  display: block;\n"
        "  margin-top: 0.75em;\n"
        "  font-size: 0.92em;\n"
        "  line-height: 1.35;\n"
        "}\n"
        "\n"
        ".titlepage .translator {\n"
        "  display: block;\n"
        "  margin-top: 0.6em;\n"
        "  font-size: 0.95em;\n"
        "}\n"
        "\n"
        ".titlepage .partial-note {\n"
        "  display: block;\n"
        "  margin-top: 2.5em;\n"
        "  font-size: 0.85em;\n"
        "  font-style: italic;\n"
        "}\n"
        "\n"
        ".imprintpage {\n"
        "  padding-top: 14%;\n"
        "  max-width: 34em;\n"
        "  margin: 0;\n"
        "  text-align: left;\n"
        "  color: #202027;\n"
        "}\n"
        "\n"
        ".imprintpage h1 {\n"
        "  margin: 0 0 2.4em;\n"
        "  font-size: 1.05em;\n"
        "  font-weight: normal;\n"
        "  letter-spacing: 0.08em;\n"
        "  text-transform: uppercase;\n"
        "  color: #56515f;\n"
        "}\n"
        "\n"
        ".imprintpage p {\n"
        "  text-align: left;\n"
        "  text-indent: 0;\n"
        "  margin: 0 0 0.95em;\n"
        "  line-height: 1.38;\n"
        "}\n"
        "\n"
        ".imprintpage .imprint-title {\n"
        "  margin-bottom: 0.45em;\n"
        "  font-size: 1.55em;\n"
        "  line-height: 1.18;\n"
        "}\n"
        "\n"
        ".imprintpage .imprint-author {\n"
        "  margin-bottom: 1.8em;\n"
        "  font-size: 1.08em;\n"
        "}\n"
        "\n"
        ".imprintpage .imprint-subtitle {\n"
        "  margin-bottom: 2em;\n"
        "  font-size: 0.98em;\n"
        "  line-height: 1.35;\n"
        "}\n"
        "\n"
        ".imprintpage .imprint-meta {\n"
        "  font-size: 0.88em;\n"
        "}\n"
        "\n"
        ".imprintpage .imprint-rights {\n"
        "  margin-top: 1.25em;\n"
        "  font-size: 0.82em;\n"
        "  line-height: 1.34;\n"
        "}\n"
        "\n"
        "blockquote {\n"
        "  margin: 1.1em 0 1.1em 0.8em;\n"
        "  padding: 0;\n"
        "  font-size: 0.88em;\n"
        "  font-style: italic;\n"
        "  line-height: 1.36;\n"
        "  color: #444;\n"
        "}\n"
        "\n"
        "blockquote p {\n"
        "  text-indent: 0;\n"
        "  margin: 0 0 0.4em;\n"
        "}\n"
    )
    path.write_text(css + "\n", encoding="utf-8")
    return path


def find_pandoc() -> str | None:
    pandoc = shutil.which("pandoc")
    if pandoc:
        return pandoc
    env_path = os.environ.get("PANDOC_PATH")
    if env_path:
        candidate = Path(env_path)
        if candidate.is_file():
            return str(candidate)
    guesses: list[Path] = []
    if sys.platform == "win32":
        guesses.extend([
            Path.home() / "AppData" / "Local" / "Pandoc" / "pandoc.exe",
            Path("C:/Program Files/Pandoc/pandoc.exe"),
            Path("C:/Program Files (x86)/Pandoc/pandoc.exe"),
        ])
    else:
        guesses.extend([
            Path("/usr/bin/pandoc"),
            Path("/usr/local/bin/pandoc"),
        ])
    for candidate in guesses:
        if candidate.is_file():
            return str(candidate)
    return None


def write_epub(
    path: Path,
    markdown_path: Path,
    cover_path: Path,
    meta: dict[str, Any],
) -> None:
    pandoc = find_pandoc()
    if not pandoc:
        raise RuntimeError(
            "Pandoc wurde nicht gefunden; EPUB-Export nicht moeglich. "
            "Bitte installieren Sie Pandoc oder setzen Sie PANDOC_PATH auf den Pfad zu pandoc.exe."
        )
    path.parent.mkdir(parents=True, exist_ok=True)
    css_path = write_epub_css(markdown_path.parent / "epub.css")
    cmd = [
        pandoc,
        str(markdown_path),
        "-o",
        str(path),
        "--toc",
        "--toc-depth=1",
        "--split-level=1",
        "--epub-chapter-level=1",
        "--epub-title-page=false",
        f"--epub-cover-image={cover_path}",
        f"--css={css_path}",
        f"--metadata=lang:{meta.get('language', 'de-DE')}",
    ]
    if meta.get("description"):
        cmd.append(f"--metadata=description:{meta['description']}")
    if meta.get("publisher"):
        cmd.append(f"--metadata=publisher:{meta['publisher']}")
    if meta.get("rights"):
        cmd.append(f"--metadata=rights:{meta['rights']}")
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=REPO_ROOT)
    if result.returncode != 0:
        raise RuntimeError(
            "Pandoc EPUB-Export fehlgeschlagen:\n"
            + result.stdout
            + result.stderr
        )


def check_epub(path: Path) -> list[str]:
    required = ["mimetype", "META-INF/container.xml"]
    missing: list[str] = []
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        xhtml_entries = [
            name for name in names
            if name.lower().endswith((".xhtml", ".html", ".htm"))
        ]
        for name in sorted(xhtml_entries):
            try:
                root = ET.fromstring(zf.read(name))
            except ET.ParseError as exc:
                missing.append(f"{name}: XHTML-Parsefehler: {exc}")
                continue
            for image in root.iter():
                if not str(image.tag).endswith("img"):
                    continue
                if image.get("src") is None:
                    missing.append(f"{name}: img ohne src")
                if image.get("alt") is None:
                    missing.append(f"{name}: img ohne alt")
    missing.extend(item for item in required if item not in names)
    if not any(name.endswith(".opf") for name in names):
        missing.append("*.opf")
    if not any("nav" in name.lower() for name in names):
        missing.append("nav")
    if not any("cover" in name.lower() for name in names):
        missing.append("cover")
    return missing


def remove_auto_title_heading_from_epub(
    path: Path,
    title: str,
    nav_title: str | None = None,
) -> None:
    """Pandoc emits a visible H1 before the custom title page.

    `--epub-title-page=false` disables Pandoc's generated title page, but the
    first body file may still start with a metadata-derived H1. Reader apps can
    paginate that as a lonely title page before our formatted title page.
    """
    with zipfile.ZipFile(path, "r") as zf:
        entries = [(info, zf.read(info.filename)) for info in zf.infolist()]

    title_re = re.escape(title.strip())
    h1_pattern = re.compile(
        rb'(<body[^>]*>\s*<section[^>]*>\s*)'
        rb'<h1[^>]*>\s*'
        + title_re.encode("utf-8")
        + rb'\s*</h1>\s*(?=<div class="frontmatter-page titlepage")',
        flags=re.DOTALL,
    )
    nav_label = (nav_title or "Titelseite").strip() or "Titelseite"
    nav_title_pattern = re.compile(
        rb'(<nav[^>]*epub:type="toc"[^>]*>.*?<ol[^>]*>\s*'
        rb'<li[^>]*>\s*<a\b[^>]*>)'
        + re.escape(html.escape(title.strip()).encode("utf-8"))
        + rb'(</a>)',
        flags=re.DOTALL,
    )

    new_entries: list[tuple[zipfile.ZipInfo, bytes]] = []
    changed = False
    for info, data in entries:
        if info.filename.endswith(".xhtml") and b"frontmatter-page titlepage" in data:
            new_data, count = h1_pattern.subn(rb"\1", data, count=1)
            if count:
                data = new_data
                changed = True
        if info.filename.endswith("nav.xhtml"):
            new_data, count = nav_title_pattern.subn(
                rb"\1" + html.escape(nav_label).encode("utf-8") + rb"\2",
                data,
                count=1,
            )
            if count:
                data = new_data
                changed = True
        new_entries.append((info, data))

    if not changed:
        return

    tmp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(tmp, "w") as zf:
        for info, data in new_entries:
            compression = zipfile.ZIP_STORED if info.filename == "mimetype" else zipfile.ZIP_DEFLATED
            new_info = zipfile.ZipInfo(info.filename, info.date_time)
            new_info.comment = info.comment
            new_info.extra = info.extra
            new_info.internal_attr = info.internal_attr
            new_info.external_attr = info.external_attr
            new_info.create_system = info.create_system
            zf.writestr(new_info, data, compress_type=compression)
    tmp.replace(path)


def rel_label(path: Path) -> str:
    try:
        return str(path.relative_to(REPO_ROOT))
    except ValueError:
        return str(path)


def output_basename(
    meta: dict[str, Any],
    style: str,
    scope: str,
    chapter: str | None,
) -> str:
    title = sanitize_filename(str(meta.get("title", "export")))
    if scope == "chapter":
        return f"chapter-{chapter}-{title}-{style}"
    return f"book-{title}-{style}"


def main() -> int:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        if hasattr(sys.stdout, "buffer"):
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8",
                                          errors="replace")

    ap = argparse.ArgumentParser(description="Exportiert DOCX/EPUB/PDF aus DE-Szenen.")
    ap.add_argument("--book", default=None, help="Buch-ID")
    ap.add_argument("--style", required=True, help="Style-Profil/Output-Ordner")
    ap.add_argument("--scope", choices=["chapter", "book"], required=True)
    ap.add_argument("--chapter", default=None, help="Kapitel-ID bei scope=chapter")
    ap.add_argument("--format", choices=EXPORT_FORMATS, default="all")
    ap.add_argument("--allow-partial", action="store_true")
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    book = find_book(args.book)
    output_root = book_output_root(REPO_ROOT, book)
    meta = load_export_config(book)
    result = collect_export(
        output_root=output_root,
        style=args.style,
        scope=args.scope,
        chapter_id=args.chapter,
        allow_partial=args.allow_partial,
        source_lang=str(book.get("source_lang") or "ru"),
    )
    print(f"=== Export: {meta.get('title', book.get('title'))} ===")
    print(f"Scope: {args.scope}")
    print(f"Style: {args.style}")
    print(f"Format: {args.format}")
    source_lang_label = str(book.get("source_lang") or "ru").upper()
    if result.missing_by_chapter:
        print("Fehlende Szenen:")
        for cid, nums in result.missing_by_chapter.items():
            print(f"  {cid}: " + ", ".join(f"{num:02d}" for num in nums))
    if result.missing_by_chapter and not args.allow_partial:
        print("ABBRUCH: unvollstaendig. Nutze --allow-partial fuer Teil-Export.")
        return 2
    if not result.chapters:
        print("ABBRUCH: keine exportierbaren Szenen gefunden.")
        return 2

    for chapter in result.chapters:
        print(
            f"Kapitel {chapter.chapter_id}: "
            f"{len(chapter.scenes)} Szenen exportierbar "
            f"({source_lang_label}={chapter.ru_count}, DE={chapter.de_count})"
        )
    illustrations = collect_illustrations(result.chapters, meta)
    if illustrations_enabled(meta):
        print(f"Illustrationen: {len(illustrations)} gefunden")
        for item in illustrations:
            print(f"  {illustration_label(item)}: {rel_label(item.path)}")
    if args.dry_run:
        print("(dry-run: keine Dateien geschrieben)")
        return 0

    dirs = export_dirs(book_exports_root(REPO_ROOT, book), args.style, args.scope)
    for path in dirs.values():
        path.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    base = output_basename(meta, args.style, args.scope, args.chapter)
    base = f"{base}-{timestamp}"
    scope_label = (
        f"Kapitel {args.chapter}" if args.scope == "chapter" else "Gesamtes Buch"
    )
    title = document_title(meta, args.scope, result.chapters[0])
    cover_path = prepare_cover(
        dirs["work"] / base,
        title=title,
        author=str(meta.get("author", book.get("author", ""))),
        style=args.style,
        scope_label=scope_label,
        meta=meta,
    )
    md_path = dirs["work"] / f"{base}.md"
    cover_ref = markdown_image_path(cover_path, md_path)
    md_text = render_export_markdown(
        result.chapters,
        meta,
        args.style,
        args.scope,
        result.partial,
        cover_ref,
        markdown_path=md_path,
        illustrations=illustrations,
    )
    md_path.write_text(md_text, encoding="utf-8")

    outputs: list[Path] = []
    if args.format in ("docx", "all"):
        docx_path = dirs["docx"] / f"{base}.docx"
        write_docx(
            docx_path, result.chapters, meta, args.style, args.scope,
            result.partial, cover_path, illustrations,
        )
        outputs.append(docx_path)
    if args.format in ("epub", "all"):
        epub_path = dirs["epub"] / f"{base}.epub"
        write_epub(epub_path, md_path, cover_path, meta)
        remove_auto_title_heading_from_epub(
            epub_path,
            document_title(meta, args.scope, result.chapters[0] if result.chapters else None),
            str(front_matter_config(meta).get("title_heading") or "Titelseite"),
        )
        missing = check_epub(epub_path)
        if missing:
            raise RuntimeError(
                "EPUB-Sanity-Check fehlgeschlagen: " + ", ".join(missing)
            )
        outputs.append(epub_path)
    pdf_html_path: Path | None = None
    pdf_css_path: Path | None = None
    if args.format == "pdf":
        pdf_path = dirs["pdf"] / f"{base}.pdf"
        pdf_html_path = dirs["work"] / f"{base}.pdf.html"
        pdf_css_path = write_pdf_css(dirs["work"] / "book-print.css")
        pdf_html = render_pdf_html(
            result.chapters,
            meta,
            args.style,
            args.scope,
            result.partial,
            cover_path,
            pdf_html_path,
            pdf_css_path,
            illustrations,
        )
        pdf_html_path.write_text(pdf_html, encoding="utf-8")
        write_pdf(pdf_path, pdf_html_path)
        outputs.append(pdf_path)
    manifest_path = dirs["manifests"] / f"{base}.json"
    manifest = {
        "book_id": book["id"],
        "book_title": meta.get("title", book.get("title")),
        "style": args.style,
        "scope": args.scope,
        "chapter": args.chapter if args.scope == "chapter" else None,
        "format": args.format,
        "partial": result.partial,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "chapters": [
            {
                "id": chapter.chapter_id,
                "scenes": len(chapter.scenes),
                "missing": chapter.missing,
            }
            for chapter in result.chapters
        ],
        "outputs": [rel_label(path) for path in outputs],
        "work_markdown": rel_label(md_path),
        "work_pdf_html": rel_label(pdf_html_path) if pdf_html_path else None,
        "work_pdf_css": rel_label(pdf_css_path) if pdf_css_path else None,
        "cover": rel_label(cover_path),
        "illustrations": [
            {
                "kind": item.kind,
                "chapter": item.chapter_id,
                "scene": item.scene_number,
                "path": rel_label(item.path),
            }
            for item in illustrations
        ],
    }
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print("Geschrieben:")
    for path in outputs:
        print(f"  {rel_label(path)}")
    print(f"Arbeitsdatei: {rel_label(md_path)}")
    print(f"Manifest: {rel_label(manifest_path)}")
    cover_label = rel_label(cover_path)
    print(f"Cover: {cover_label}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
